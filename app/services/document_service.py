"""
Document processing service for file upload, text extraction, and chunking
"""
import os
import uuid
import aiofiles
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, BinaryIO
from fastapi import UploadFile, HTTPException, status
from sqlalchemy.orm import Session
from sqlalchemy import and_

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument

from app.models.document import Document, DocumentChunk
from app.models.user import User
from app.services.embedding_service import EmbeddingService
from app.services.vector_service import QdrantVectorService
from app.config import settings

logger = logging.getLogger(__name__)


class DocumentService:
    """
    Service for handling document upload, processing, and management
    """
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert to bytes
        self.allowed_types = settings.allowed_file_types
        
        # Initialize services
        self.embedding_service = EmbeddingService()
        self.vector_service = QdrantVectorService()
        
        # Ensure upload directory exists
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def _validate_file(self, file: UploadFile) -> bool:
        """
        Validate uploaded file
        """
        # Check file size
        if hasattr(file, 'size') and file.size > self.max_file_size:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size exceeds maximum allowed size of {settings.max_file_size_mb}MB"
            )
        
        # Check file type
        if file.filename:
            file_ext = file.filename.split('.')[-1].lower()
            if file_ext not in self.allowed_types:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File type '{file_ext}' not allowed. Allowed types: {', '.join(self.allowed_types)}"
                )
        
        return True
    
    async def upload_document(
        self,
        db: Session,
        user_id: str,
        file: UploadFile,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Upload and store document file
        """
        # Validate file
        self._validate_file(file)
        
        # Generate unique filename
        unique_filename = f"{user_id}_{str(uuid.uuid4())}.{file.filename.split('.')[-1].lower()}"
        file_path = self.upload_dir / unique_filename
        
        try:
            # Save file to disk
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            # Create document record
            document = Document(
                user_id=user_id,
                filename=unique_filename,
                original_filename=file.filename,
                content_type=file.content_type,
                file_size=len(content),
                file_path=str(file_path),
                status="uploaded",
                doc_metadata=metadata or {}
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            logger.info(f"Document uploaded: {file.filename} for user {user_id}")
            return document
            
        except Exception as e:
            # Clean up file if database operation fails
            if file_path.exists():
                file_path.unlink()
            logger.error(f"Document upload failed: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to upload document"
            )
    
    def extract_text_from_file(self, file_path: str, content_type: str) -> str:
        """
        Extract text content from various file types
        """
        try:
            if content_type == "application/pdf" or file_path.endswith('.pdf'):
                return self._extract_from_pdf(file_path)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith('.docx'):
                return self._extract_from_docx(file_path)
            elif content_type == "text/plain" or file_path.endswith('.txt'):
                return self._extract_from_txt(file_path)
            else:
                # Try as plain text
                return self._extract_from_txt(file_path)
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            logger.info(f"PDF extraction start: pages={len(pdf_reader.pages)} file={file_path}")
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                else:
                    logger.warning(f"PDF page had no extractable text: page={page_num+1}")
        
        combined = "\n\n".join(text_content)
        logger.info(f"PDF extraction done: total_chars={len(combined)} file={file_path}")
        return combined
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return "\n\n".join(text_content)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    async def process_document(
        self,
        db: Session,
        document_id: str,
        user_id: str
    ) -> bool:
        """
        Process document: extract text, create chunks, generate embeddings
        """
        # Validate user exists
        user = db.query(User).filter(
            and_(
                User.id == user_id,
                User.is_active == True
            )
        ).first()
        
        if not user:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="User not found or inactive"
            )
        
        # Get document
        document = db.query(Document).filter(
            and_(
                Document.id == document_id,
                Document.user_id == user_id
            )
        ).first()
        
        if not document:
            logger.error(f"Document {document_id} not found for user {user_id}")
            return False
        
        try:
            # Update status
            document.status = "processing"
            db.commit()
            
            # Extract text content
            text_content = self.extract_text_from_file(
                document.file_path, 
                document.content_type
            )
            logger.info(
                f"Doc extract stats: id={document.id} user={user_id} content_type={document.content_type} chars={len(text_content) if text_content else 0}"
            )
            
            if not text_content.strip():
                document.status = "failed"
                db.commit()
                logger.error(f"No text content extracted from document {document_id}")
                return False
            
            # Generate basic metadata
            word_count = len(text_content.split())
            document.word_count = word_count
            logger.info(f"Doc words: id={document.id} words={word_count}")
            
            # Create text chunks
            chunks = self.embedding_service.chunk_text_for_embedding(
                text_content,
                max_chunk_size=512,
                overlap_size=50
            )
            logger.info(
                f"Chunking done: id={document.id} chunks={len(chunks)} first_chunk_chars={len(chunks[0]['text']) if chunks else 0}"
            )
            
            # Generate embeddings for chunks
            embedded_chunks = await self.embedding_service.embed_document_chunks(chunks)
            # Log embedding lengths for first few chunks
            try:
                emb_dims = [len(c.get("embedding", []) or []) for c in embedded_chunks[:3]]
            except Exception:
                emb_dims = []
            logger.info(f"Embeddings generated: id={document.id} chunks_embedded={len(embedded_chunks)} sample_dims={emb_dims}")
            
            # Store chunks in database and vector store
            vector_docs = []
            chunk_records = []
            
            for chunk_data in embedded_chunks:
                # Pre-generate a stable UUID for this chunk id
                chunk_uuid = uuid.uuid4()
                # Create database record
                chunk_record = DocumentChunk(
                    id=chunk_uuid,
                    document_id=document.id,
                    user_id=user_id,
                    chunk_index=chunk_data["chunk_index"],
                    text_content=chunk_data["text"],
                    chunk_size=chunk_data["chunk_size"],
                    start_char=chunk_data["start_char"],
                    end_char=chunk_data["end_char"],
                    vector_id=str(uuid.uuid4()),
                    embedding_model=chunk_data["embedding_model"],
                    embedding_dimension=chunk_data["embedding_dimension"]
                )
                
                chunk_records.append(chunk_record)
                
                # Prepare for vector store
                vector_doc = {
                    "document_id": str(document.id),
                    "chunk_id": str(chunk_uuid),
                    "text": chunk_data["text"],
                    "embedding": chunk_data["embedding"],
                    "source": document.original_filename,  # Add source field
                    "page_number": chunk_data.get("page_number"),
                    "metadata": {
                        "filename": document.original_filename,
                        "content_type": document.content_type,
                        "chunk_index": chunk_data["chunk_index"],
                        "start_char": chunk_data["start_char"],
                        "end_char": chunk_data["end_char"]
                    }
                }
                vector_docs.append(vector_doc)
            logger.info(
                f"Vector prep: id={document.id} vectors={len(vector_docs)} sample_vec_len={len(vector_docs[0]['embedding']) if vector_docs else 0}"
            )
            
            # Store chunks in database
            db.add_all(chunk_records)
            
            # Store embeddings in vector database
            success = await self.vector_service.add_documents(
                user_id=user_id,
                documents=vector_docs
            )
            
            if success:
                # Update document status
                document.status = "processed"
                document.total_chunks = len(chunks)
                document.processed_chunks = len(chunks)
                document.collection_name = self.vector_service.default_collection
                document.embedding_model = embedded_chunks[0]["embedding_model"]
                
                db.commit()
                
                logger.info(f"Document {document_id} processed successfully with {len(chunks)} chunks (user={user_id})")
                return True
            else:
                document.status = "failed"
                db.commit()
                logger.error(f"Vector upsert failed: id={document.id} user={user_id}")
                return False
                
        except Exception as e:
            document.status = "failed"
            db.commit()
            logger.error(f"Document processing failed for {document_id}: {e}")
            return False
    
    async def delete_document(
        self,
        db: Session,
        document_id: str,
        user_id: str
    ) -> bool:
        """
        Delete document and all associated data
        """
        try:
            # Get document
            document = db.query(Document).filter(
                and_(
                    Document.id == document_id,
                    Document.user_id == user_id
                )
            ).first()
            
            if not document:
                return False
            
            # Delete from vector store
            await self.vector_service.delete_document(user_id, str(document_id))
            
            # Delete file from disk
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete from database (cascades to chunks)
            db.delete(document)
            db.commit()
            
            logger.info(f"Document {document_id} deleted successfully")
            return True
            
        except Exception as e:
            logger.error(f"Document deletion failed for {document_id}: {e}")
            return False
    
    def get_document(
        self,
        db: Session,
        document_id: str,
        user_id: str
    ) -> Optional[Document]:
        """
        Get document with user validation
        """
        return db.query(Document).filter(
            and_(
                Document.id == document_id,
                Document.user_id == user_id
            )
        ).first()
    
    def list_documents(
        self,
        db: Session,
        user_id: str,
        skip: int = 0,
        limit: int = 100,
        status_filter: Optional[str] = None
    ) -> List[Document]:
        """
        List documents for a user with pagination and filtering
        """
        # Base query with user filter
        query = db.query(Document).filter(Document.user_id == user_id)
        
        if status_filter:
            query = query.filter(Document.status == status_filter)
        
        return query.offset(skip).limit(limit).all()